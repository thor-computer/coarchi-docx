import argparse
import csv
import os
import re
import tempfile
from docx import Document


def convert_md_to_docx(md_path, docx_path):
    """Convert a markdown file to docx format."""
    # Read the markdown file
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create a new Document
    doc = Document()
    
    for line in lines:
        line = line.rstrip('\n')
        
        if not line.strip():
            # Empty line
            doc.add_paragraph()
        elif line.startswith('# '):
            # Heading 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            # Heading 2
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # Heading 3
            doc.add_heading(line[4:], level=3)
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Save the document
    doc.save(docx_path)
    print(f"Converted markdown to docx: {docx_path}")


def extract_markdown_from_docx(docx_path):
    """Extract text from docx and convert to markdown format."""
    doc = Document(docx_path)
    
    # Extract text from paragraphs, converting heading styles to markdown
    lines = []
    for para in doc.paragraphs:
        text = para.text
        # Check if paragraph is a heading
        if para.style.name.startswith('Heading'):
            # Extract heading level (e.g., 'Heading 1' -> 1)
            try:
                level = int(para.style.name.split()[-1])
                # Convert to markdown format
                text = '#' * level + ' ' + text
            except (ValueError, IndexError):
                pass
        lines.append(text)
    
    return '\n'.join(lines)


def parse_docx_file(docx_path):
    """Parse the .docx file and extract structured data."""
    content = extract_markdown_from_docx(docx_path)
    return parse_content(content)


def parse_txt_file(txt_path):
    """Parse the .txt file and extract structured data."""
    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    return parse_content(content)


def parse_content(content):
    """Parse text content and extract structured data."""
    
    rows = []
    
    # Split content by sections (marked by # headers)
    sections = re.split(r'\n(?=# )', content.strip())
    
    for section in sections:
        if not section.strip():
            continue
        
        lines = section.strip().split('\n')
        
        # Initialize data structure
        data = {
            'Application': '',
            'Box Location': 'External',
            'Box Type': 'Integration',
            'Group Type': 'Topic',
            'Group Title': '',
            'Box ID': '',
            'Box Title': '',
            'Arrow Direction': 'To',
            'Arrow Description': ''
        }
        
        # Parse the section
        for line in lines:
            line = line.strip()
            
            # Application (# header)
            if line.startswith('# '):
                data['Application'] = line[2:].strip()
                data['Group Title'] = 'Integrations'  # Default value
            
            # Box ID (## header)
            elif line.startswith('## '):
                data['Box ID'] = line[3:].strip()
            
            # Function becomes Box Title
            elif line.startswith('Function:'):
                data['Box Title'] = line.split(':', 1)[1].strip()
            
            # Business purpose becomes Arrow Description
            elif line.startswith('Business purpose:'):
                data['Arrow Description'] = line.split(':', 1)[1].strip()
        
        # Only add if we have meaningful data
        if data['Application'] and data['Box ID']:
            rows.append(data)
    
    return rows


def write_csv(data, output_path):
    """Write data to CSV file matching the format of csv.csv."""
    # Define the CSV headers matching the original format
    headers = [
        'Application',
        'Box Location',
        'Box Type',
        'Group Type',
        'Group Title',
        'Box ID',
        'Box Title',
        'Arrow Direction',
        'Arrow Description'
    ]
    
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=headers, delimiter=';')
        writer.writeheader()
        
        for row in data:
            # Map to match the header with space
            csv_row = {
                'Application': row['Application'],
                'Box Location': row['Box Location'],
                'Box Type': row['Box Type'],
                'Group Type': row['Group Type'],
                'Group Title': row['Group Title'],
                'Box ID': row['Box ID'],
                'Box Title': row['Box Title'],
                'Arrow Direction': row['Arrow Direction'],
                'Arrow Description': row['Arrow Description']
            }
            writer.writerow(csv_row)


def main():
    parser = argparse.ArgumentParser(description='Convert .txt, .md, or .docx file to .csv format')
    parser.add_argument('--document', required=True, help='Path to the input .txt, .md, or .docx file')
    
    args = parser.parse_args()
    
    # Validate input file exists
    if not os.path.exists(args.document):
        print(f"Error: Input file '{args.document}' not found.")
        return 1
    
    # Parse the input file based on extension
    print(f"Processing {args.document}...")
    file_ext = os.path.splitext(args.document)[1].lower()
    
    # Create output directory if it doesn't exist
    output_dir = 'output'
    os.makedirs(output_dir, exist_ok=True)
    
    docx_to_extract = None
    
    if file_ext == '.md':
        # Convert markdown to docx first, then parse
        temp_docx = os.path.join(output_dir, 'docx.docx')
        convert_md_to_docx(args.document, temp_docx)
        print(f"Parsing {temp_docx}...")
        data = parse_docx_file(temp_docx)
        docx_to_extract = temp_docx
    elif file_ext == '.docx':
        print(f"Parsing {args.document}...")
        data = parse_docx_file(args.document)
        docx_to_extract = args.document
    elif file_ext == '.txt':
        print(f"Parsing {args.document}...")
        data = parse_txt_file(args.document)
    else:
        print(f"Error: Unsupported file type '{file_ext}'. Only .txt, .md, and .docx are supported.")
        return 1
    
    # Extract markdown from docx if we processed a docx file
    if docx_to_extract:
        md_output_path = os.path.join(output_dir, 'md.md')
        markdown_content = extract_markdown_from_docx(docx_to_extract)
        with open(md_output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"Extracted markdown to {md_output_path}")
    
    # Generate output filename (always csv.csv)
    output_filename = 'csv.csv'
    output_path = os.path.join(output_dir, output_filename)
    
    # Write CSV
    print(f"Writing output to {output_path}...")
    write_csv(data, output_path)
    
    print(f"Successfully converted {len(data)} record(s) to CSV.")
    print(f"Output files: {output_path}" + (f", {os.path.join(output_dir, 'md.md')}" if docx_to_extract else ""))
    return 0


if __name__ == '__main__':
    exit(main())
